#!/usr/bin/env python3
"""Build the synthetic PDF and DOCX inputs used by the live compiler eval."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[2]
SCHEMA_PATH = ROOT / "data" / "example_entreeformulier.schema.json"

INK = HexColor("#17352D")
GREEN = HexColor("#426A5A")
PALE = HexColor("#EDF4F0")
MUTED = HexColor("#61716B")
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
DOC_INK = RGBColor(0x17, 0x35, 0x2D)
DOC_GREEN = RGBColor(0x42, 0x6A, 0x5A)
DOC_MUTED = RGBColor(0x61, 0x71, 0x6B)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", type=Path, default=ROOT / "work" / "golden")
    args = parser.parse_args()
    args.out.mkdir(parents=True, exist_ok=True)

    medical_path = args.out / "medical-intake.pdf"
    school_path = args.out / "elementary-school-intake.docx"
    build_medical_pdf(medical_path)
    build_school_docx(school_path)
    audit_school_docx(school_path)
    print(json.dumps({"medical": str(medical_path), "school": str(school_path)}, indent=2))


def build_medical_pdf(output: Path) -> None:
    pdf = canvas.Canvas(str(output), pagesize=letter, pageCompression=1)
    width, height = letter
    pdf.setTitle("Riverside Family Practice - New Patient Intake")
    pdf.setAuthor("VocaForm synthetic evaluation fixture")

    pdf.setFillColor(PALE)
    pdf.roundRect(45, height - 124, width - 90, 76, 12, fill=1, stroke=0)
    pdf.setFillColor(GREEN)
    pdf.setFont("Helvetica-Bold", 9)
    pdf.drawString(62, height - 75, "NEW PATIENT INTAKE")
    pdf.setFillColor(INK)
    pdf.setFont("Helvetica-Bold", 22)
    pdf.drawString(62, height - 102, "Riverside Family Practice")
    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 9)
    pdf.drawRightString(width - 62, height - 75, "Synthetic evaluation form")

    y = height - 158
    y = section_heading(pdf, "PATIENT DETAILS", y)
    y = answer_line(pdf, "Full legal name (required):", y)
    y = answer_line(pdf, "Date of birth (required):", y, hint="MM / DD / YYYY")
    y = answer_line(pdf, "Phone number (required):", y)
    y = answer_line(pdf, "Email address (optional):", y)

    y -= 10
    y = section_heading(pdf, "VISIT DETAILS", y)
    y = answer_box(pdf, "Reason for today's visit (required):", y, 48)
    y = answer_box(pdf, "List current medications, or write none:", y, 48)

    pdf.setFillColor(INK)
    pdf.setFont("Helvetica-Bold", 10.5)
    pdf.drawString(62, y, "Do you have any known allergies? Yes / No (required)")
    checkbox(pdf, width - 174, y - 4, "Yes")
    checkbox(pdf, width - 112, y - 4, "No")
    y -= 34
    y = answer_box(pdf, "If yes, list the allergies and reactions:", y, 48)

    pdf.setStrokeColor(HexColor("#CAD8D1"))
    pdf.line(62, 49, width - 62, 49)
    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 8)
    pdf.drawString(62, 34, "This is synthetic demonstration paperwork. It contains no real patient data.")
    pdf.drawRightString(width - 62, 34, "VocaForm compiler evaluation")
    pdf.save()


def section_heading(pdf: canvas.Canvas, title: str, y: float) -> float:
    pdf.setFillColor(GREEN)
    pdf.setFont("Helvetica-Bold", 10)
    pdf.drawString(62, y, title)
    pdf.setStrokeColor(HexColor("#9CB7AA"))
    pdf.line(62, y - 7, 550, y - 7)
    return y - 29


def answer_line(pdf: canvas.Canvas, label: str, y: float, hint: str | None = None) -> float:
    pdf.setFillColor(INK)
    pdf.setFont("Helvetica-Bold", 10.5)
    pdf.drawString(62, y, label)
    pdf.setStrokeColor(HexColor("#879A91"))
    pdf.line(62, y - 18, 550, y - 18)
    if hint:
        pdf.setFillColor(MUTED)
        pdf.setFont("Helvetica", 8)
        pdf.drawRightString(550, y - 14, hint)
    return y - 42


def answer_box(pdf: canvas.Canvas, label: str, y: float, box_height: float) -> float:
    pdf.setFillColor(INK)
    pdf.setFont("Helvetica-Bold", 10.5)
    pdf.drawString(62, y, label)
    pdf.setStrokeColor(HexColor("#A9BBB2"))
    pdf.setFillColor(white)
    pdf.roundRect(62, y - box_height - 10, 488, box_height, 5, fill=1, stroke=1)
    return y - box_height - 31


def checkbox(pdf: canvas.Canvas, x: float, y: float, label: str) -> None:
    pdf.setStrokeColor(GREEN)
    pdf.rect(x, y, 11, 11, fill=0, stroke=1)
    pdf.setFillColor(INK)
    pdf.setFont("Helvetica", 9)
    pdf.drawString(x + 16, y + 1, label)


def build_school_docx(output: Path) -> None:
    schema = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
    doc = Document()
    doc.core_properties.title = "Entreeformulier Dit ben ik"
    doc.core_properties.subject = "Synthetic elementary school intake questionnaire"
    doc.core_properties.author = "VocaForm synthetic evaluation fixture"

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    configure_header_footer(section)
    add_customer_pack_title(doc)

    intro = doc.add_paragraph()
    intro.add_run(
        "Vertel ons wat uw kind nodig heeft om zich veilig, gezien en nieuwsgierig te voelen op school. "
        "Vragen met een * zijn verplicht."
    )

    for source_section in schema["sections"]:
        heading = doc.add_paragraph(source_section["title"], style="Heading 1")
        heading.paragraph_format.keep_with_next = True
        for field in source_section["fields"]:
            add_question(doc, field)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Dit is synthetisch demonstratiepapier. Het bevat geen echte gegevens van kinderen of gezinnen.")
    set_run_font(run, "Calibri", 8.5, DOC_MUTED, italic=True)
    doc.save(output)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, DOC_INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    configure_heading(doc.styles["Heading 1"], 16, BLUE, 18, 10)
    configure_heading(doc.styles["Heading 2"], 13, BLUE, 14, 7)
    configure_heading(doc.styles["Heading 3"], 12, DARK_BLUE, 10, 5)

    question = doc.styles.add_style("Form Question", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(question, "Calibri", 10.5, DOC_INK, bold=True)
    question.paragraph_format.space_before = Pt(3)
    question.paragraph_format.space_after = Pt(2)
    question.paragraph_format.line_spacing = 1.1

    answer = doc.styles.add_style("Form Answer Line", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(answer, "Calibri", 9, DOC_MUTED)
    answer.paragraph_format.space_before = Pt(0)
    answer.paragraph_format.space_after = Pt(6)
    answer.paragraph_format.line_spacing = 1.0


def configure_heading(style, size: float, color: RGBColor, before: float, after: float) -> None:
    set_style_font(style, "Calibri", size, color, bold=True)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_with_next = True


def configure_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    left = header.add_run("DIT BEN IK")
    set_run_font(left, "Calibri", 8.5, DOC_GREEN, bold=True)
    right = header.add_run("   |   Synthetisch schoolformulier")
    set_run_font(right, "Calibri", 8.5, DOC_MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    label = footer.add_run("VocaForm evaluatie   |   pagina ")
    set_run_font(label, "Calibri", 8, DOC_MUTED)
    add_page_field(footer)


def add_customer_pack_title(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(2)
    kicker.paragraph_format.space_after = Pt(0)
    run = kicker.add_run("SCHOOL STARTVRAGENLIJST")
    set_run_font(run, "Calibri", 9, DOC_GREEN, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("Dit ben ik")
    set_run_font(run, "Calibri", 28, DOC_INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Entreeformulier over ontwikkeling, spel, taal en het gezin")
    set_run_font(run, "Calibri", 12.5, DOC_MUTED)


def add_question(doc: Document, field: dict) -> None:
    question = doc.add_paragraph(style="Form Question")
    question.paragraph_format.keep_with_next = True
    question.add_run(field.get("render_anchor") or field["label"])
    if field.get("required"):
        marker = question.add_run("  *")
        marker.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)

    answer = doc.add_paragraph(style="Form Answer Line")
    answer.paragraph_format.keep_together = True
    answer.add_run("Antwoord")
    add_bottom_border(answer, "A9BBB2")


def add_bottom_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, "Calibri", 8, DOC_MUTED)


def set_style_font(style, name: str, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def set_run_font(
    run,
    name: str,
    size: float,
    color: RGBColor,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def audit_school_docx(path: Path) -> None:
    doc = Document(path)
    section = doc.sections[0]
    assert round(section.page_width.inches, 3) == 8.5
    assert round(section.page_height.inches, 3) == 11
    assert all(
        round(value.inches, 3) == 1
        for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)
    )
    assert round(section.header_distance.inches, 3) == 0.492
    assert round(section.footer_distance.inches, 3) == 0.492
    normal = doc.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size.pt == 11
    assert normal.paragraph_format.space_after.pt == 6
    assert normal.paragraph_format.line_spacing == 1.25
    expected_questions = sum(len(section["fields"]) for section in json.loads(SCHEMA_PATH.read_text())["sections"])
    actual_questions = sum(1 for paragraph in doc.paragraphs if paragraph.style.name == "Form Question")
    assert actual_questions == expected_questions


if __name__ == "__main__":
    main()
